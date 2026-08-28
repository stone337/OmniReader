import os
import re
import tempfile
import unicodedata
from docx import Document
from ebooklib import epub
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
from PIL import Image

class ReaderEngine:
    """高保真英语文档解析器与排版引擎。
    解析 DOCX, EPUB, PDF, TXT 文件，并输出易于在 tk.Text 中进行富文本渲染的结构化段落/Runs。
    """
    def __init__(self):
        pass

    def parse_file(self, file_path):
        """解析文件入口，返回结构化的段落列表。
        结构：
        [
            {
                'type': 'paragraph',
                'alignment': 'left'|'center'|'right'|'justify',
                'indent': True|False,  # 是否首行缩进
                'runs': [
                    {
                        'text': 'text content',
                        'bold': True|False,
                        'italic': True|False,
                        'font_name': 'Arial'|None,
                        'font_size': 14|None,
                        'color': '#1e293b'|None,
                        'is_image': False
                    },
                    ...
                ]
            },
            {
                'type': 'image',
                'image_path': 'temp_path.png'
            },
            ...
        ]
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".epub":
            return self._parse_epub(file_path)
        elif ext == ".pdf":
            return self._parse_pdf_reflow(file_path)
        elif ext == ".txt":
            return self._parse_txt(file_path)
        return []

    def _parse_docx(self, file_path):
        """深度扫描 DOCX runs，保持字体、颜色、字号与加粗等状态"""
        paragraphs_data = []
        try:
            doc = Document(file_path)
            for p in doc.paragraphs:
                text_p = p.text.strip()
                if not text_p:
                    continue
                
                # 判定居中、靠右等排版
                alignment = 'left'
                align_val = p.alignment
                if align_val == 1:  # CENTER
                    alignment = 'center'
                elif align_val == 2:  # RIGHT
                    alignment = 'right'
                elif align_val == 3:  # JUSTIFY
                    alignment = 'justify'

                # 提取 runs
                runs_data = []
                for r in p.runs:
                    txt = r.text
                    if not txt:
                        continue
                    # 字符清洗物理规范化
                    txt = unicodedata.normalize('NFKC', txt).replace('\x0b', ' ').replace('\xa0', ' ')
                    
                    # 获取颜色 Hex 字符串
                    color_hex = None
                    if r.font.color and r.font.color.rgb:
                        rgb = r.font.color.rgb
                        color_hex = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"

                    # 获取字号 pt
                    font_size = None
                    if r.font.size:
                        font_size = r.font.size.pt

                    runs_data.append({
                        'text': txt,
                        'bold': bool(r.bold),
                        'italic': bool(r.italic),
                        'font_name': r.font.name,
                        'font_size': font_size,
                        'color': color_hex,
                        'is_image': False
                    })

                if runs_data:
                    paragraphs_data.append({
                        'type': 'paragraph',
                        'alignment': alignment,
                        'indent': alignment == 'left' and not p.style.name.startswith('Heading'),
                        'runs': runs_data
                    })

            # 处理文档中的表格（将单元格作为单独段落导入）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt_cell = cell.text.strip()
                        if txt_cell:
                            txt_cell = unicodedata.normalize('NFKC', txt_cell).replace('\x0b', ' ').replace('\xa0', ' ')
                            paragraphs_data.append({
                                'type': 'paragraph',
                                'alignment': 'left',
                                'indent': False,
                                'runs': [{
                                    'text': txt_cell,
                                    'bold': False,
                                    'italic': False,
                                    'font_name': None,
                                    'font_size': None,
                                    'color': '#475569',  # 表格内容使用优雅的铅灰色
                                    'is_image': False
                                }]
                            })
        except Exception as e:
            print(f"[DOCX解析报错] {e}")
        return paragraphs_data

    def _parse_epub(self, file_path):
        """EPUB HTML 深度无损解析，提取章节 HTML 标签与内置插图"""
        paragraphs_data = []
        try:
            book = epub.read_epub(file_path)
            temp_dir = tempfile.gettempdir()
            
            # 提取所有图片媒体
            images_map = {}
            for item in book.get_items():
                if item.get_type() == 9: # Image
                    name = os.path.basename(item.get_name())
                    try:
                        img_data = item.get_content()
                        t_path = os.path.join(temp_dir, f"epub_img_{uuid.uuid4().hex}_{name}")
                        with open(t_path, "wb") as f_img:
                            f_img.write(img_data)
                        images_map[item.get_name()] = t_path
                    except Exception as ex:
                        print(f"[解密EPUB图片报错] {ex}")

            # 解析文本 HTML 章节
            for item in book.get_items_of_type(9): # HTML Document
                html_content = item.get_content()
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # 遍历 body 内部的主要块级标签
                body = soup.find('body')
                if not body:
                    body = soup
                
                for element in body.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'div', 'img']):
                    if element.name == 'img':
                        src = element.get('src', '')
                        # 查找对应的图片
                        matched_path = None
                        for k, v in images_map.items():
                            if os.path.basename(k) == os.path.basename(src):
                                matched_path = v
                                break
                        if matched_path and os.path.exists(matched_path):
                            paragraphs_data.append({
                                'type': 'image',
                                'image_path': matched_path
                            })
                        continue

                    # 处理文本标签
                    text_p = element.get_text().strip()
                    if not text_p:
                        continue
                    
                    # 防止父级标签 div 和子级标签 p 重复抓取
                    if element.name == 'div' and element.find(['p', 'div']):
                        continue

                    # 判定标题样式
                    is_heading = element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']
                    alignment = 'center' if is_heading else 'left'
                    
                    # 扫描子标签获取 runs (粗体、斜体、色彩样式)
                    runs_data = []
                    for child in element.children:
                        if child.name is None: # 纯文本节点
                            txt = str(child).strip()
                            if txt:
                                runs_data.append({
                                    'text': txt + ' ',
                                    'bold': is_heading,
                                    'italic': False,
                                    'font_name': None,
                                    'font_size': 18 if is_heading else None,
                                    'color': '#0f172a' if is_heading else None,
                                    'is_image': False
                                })
                        else: # 子标签节点，如 <strong>, <em>, <span>
                            sub_text = child.get_text()
                            if not sub_text:
                                continue
                            
                            is_bold = is_heading or child.name in ['strong', 'b']
                            is_italic = child.name in ['em', 'i']
                            
                            # 简单解析 inline CSS 样式中的颜色
                            style = child.get('style', '')
                            color_hex = None
                            if 'color' in style:
                                c_match = re.search(r'color\s*:\s*(#[0-9a-fA-F]{3,6})', style)
                                if c_match:
                                    color_hex = c_match.group(1)

                            runs_data.append({
                                'text': sub_text,
                                'bold': is_bold,
                                'italic': is_italic,
                                'font_name': None,
                                'font_size': None,
                                'color': color_hex,
                                'is_image': False
                            })

                    if runs_data:
                        paragraphs_data.append({
                            'type': 'paragraph',
                            'alignment': alignment,
                            'indent': not is_heading,
                            'runs': runs_data
                        })
        except Exception as e:
            print(f"[EPUB解析报错] {e}")
        return paragraphs_data

    def _parse_pdf_reflow(self, file_path):
        """PDF 高精文本流式提取与句子重组重排引擎"""
        paragraphs_data = []
        try:
            reader = PdfReader(file_path)
            for page_idx, page in enumerate(reader.pages):
                raw_text = page.extract_text()
                if not raw_text:
                    continue

                # 1. 物理换行纠偏：仅将跨行连字(如 de- \n velopment)物理拼接复原
                raw_text = re.sub(r'-\s*\n\s*', '', raw_text)
                
                # 2. 智能合并破碎行：段落内部被 PDF 强制打断的换行，重新融合成流动段落
                lines = [line.strip() for line in raw_text.split('\n')]
                merged_paragraphs = []
                current_p = []

                for line in lines:
                    if not line:
                        if current_p:
                            merged_paragraphs.append(" ".join(current_p))
                            current_p = []
                        continue

                    current_p.append(line)
                    # 判定行尾是否是强终止符（句号、问号、感叹号、冒号、双引号等）
                    if line[-1] in ['.', '!', '?', ':', '"', ';', '。', '！', '？', '”']:
                        merged_paragraphs.append(" ".join(current_p))
                        current_p = []

                if current_p:
                    merged_paragraphs.append(" ".join(current_p))

                # 3. 将合并好的流式段落输出为结构化 Tags 格式
                for p_text in merged_paragraphs:
                    p_text_clean = p_text.strip()
                    if not p_text_clean:
                        continue
                    
                    paragraphs_data.append({
                        'type': 'paragraph',
                        'alignment': 'justify',
                        'indent': True,
                        'runs': [{
                            'text': p_text_clean,
                            'bold': False,
                            'italic': False,
                            'font_name': None,
                            'font_size': None,
                            'color': None,
                            'is_image': False
                        }]
                    })
        except Exception as e:
            print(f"[PDF解析重排报错] {e}")
        return paragraphs_data

    def _parse_txt(self, file_path):
        """解析纯文本 TXT，智能识别空行，自适应流式段落"""
        paragraphs_data = []
        try:
            # 智能多编码协商打开
            txt_content = None
            for enc in ["utf-8", "gbk", "utf-8-sig", "latin-1"]:
                try:
                    with open(file_path, "r", encoding=enc) as f:
                        txt_content = f.read()
                    break
                except:
                    pass

            if txt_content:
                txt_content = unicodedata.normalize('NFKC', txt_content)
                lines = [line.strip() for line in txt_content.split('\n')]
                for line in lines:
                    if not line:
                        continue
                    paragraphs_data.append({
                        'type': 'paragraph',
                        'alignment': 'left',
                        'indent': True,
                        'runs': [{
                            'text': line,
                            'bold': False,
                            'italic': False,
                            'font_name': None,
                            'font_size': None,
                            'color': None,
                            'is_image': False
                        }]
                    })
        except Exception as e:
            print(f"[TXT解析报错] {e}")
        return paragraphs_data
